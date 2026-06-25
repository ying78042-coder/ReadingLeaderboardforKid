from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Reading Leaderboard File Guide.docx"


BLUE = RGBColor(0x2E, 0x74, 0xB5)
DARK_BLUE = RGBColor(0x1F, 0x4D, 0x78)
MUTED = RGBColor(0x55, 0x55, 0x55)
HEADER_FILL = "E8EEF5"
BORDER = "AEB7C2"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, bottom=80, start=120, end=120):
    tbl_pr = table._tbl.tblPr
    tbl_cell_mar = tbl_pr.find(qn("w:tblCellMar"))
    if tbl_cell_mar is None:
        tbl_cell_mar = OxmlElement("w:tblCellMar")
        tbl_pr.append(tbl_cell_mar)

    for margin_name, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tbl_cell_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tbl_cell_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), BORDER)


def set_table_indent(table, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    indent = tbl_pr.find(qn("w:tblInd"))
    if indent is None:
        indent = OxmlElement("w:tblInd")
        tbl_pr.append(indent)
    indent.set(qn("w:w"), str(indent_dxa))
    indent.set(qn("w:type"), "dxa")


def set_fixed_table_width(table, widths):
    table.autofit = False
    tbl = table._tbl
    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(0, tbl_grid)

    for child in list(tbl_grid):
        tbl_grid.remove(child)

    for width in widths:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    for row in table.rows:
        for index, width in enumerate(widths):
            cell = row.cells[index]
            cell.width = Inches(width / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")


def add_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_fixed_table_width(table, widths)
    set_table_indent(table)
    set_cell_margins(table)
    set_table_borders(table)

    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_shading(cell, HEADER_FILL)
        paragraph = cell.paragraphs[0]
        paragraph.style = "Table Text"
        run = paragraph.add_run(header)
        run.bold = True

    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            paragraph = cell.paragraphs[0]
            paragraph.style = "Table Text"
            paragraph.add_run(value)

    return table


def style_document(document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 18, 10),
        ("Heading 2", 13, BLUE, 14, 7),
        ("Heading 3", 12, DARK_BLUE, 10, 5),
    ):
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25

    table_text = styles.add_style("Table Text", 1)
    table_text.font.name = "Calibri"
    table_text.font.size = Pt(9.5)
    table_text.paragraph_format.space_before = Pt(0)
    table_text.paragraph_format.space_after = Pt(0)
    table_text.paragraph_format.line_spacing = 1.15

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.style = normal
    run = footer.add_run("Reading Leaderboard File Guide")
    run.font.color.rgb = MUTED
    run.font.size = Pt(9)


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.375)
    paragraph.paragraph_format.first_line_indent = Inches(-0.188)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.25
    paragraph.add_run(text)


def build():
    document = Document()
    style_document(document)

    title = document.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("Reading Leaderboard File Guide")
    run.font.name = "Calibri"
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = BLUE

    subtitle = document.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(14)
    sub_run = subtitle.add_run("How the local file-backed kid reading leaderboard is organized")
    sub_run.font.color.rgb = MUTED

    document.add_heading("Overview", level=1)
    document.add_paragraph(
        "This app is a small local web application. The browser shows the leaderboard, "
        "the Python server reads and writes the saved records file, and the JSON file "
        "stores every reader record."
    )
    add_bullet(document, "Open the app through http://127.0.0.1:5173/ so the page can talk to the server.")
    add_bullet(document, "Each new reader is appended to data/readers.json instead of replacing earlier records.")
    add_bullet(document, "The Remove button deletes only the selected reader record from the JSON file.")

    document.add_heading("File Map", level=1)
    add_table(
        document,
        ["File", "Purpose", "Key responsibilities"],
        [
            (
                "index.html",
                "Defines the visible page structure.",
                "Builds the reader form, goal inputs, progress cards, and leaderboard list container.",
            ),
            (
                "styles.css",
                "Controls the visual design and responsive layout.",
                "Styles the two-panel app, form fields, leaderboard rows, progress meters, empty state, and remove buttons.",
            ),
            (
                "app.js",
                "Runs the browser-side behavior.",
                "Loads records from the server, renders all readers, appends new readers, removes one reader, clears all readers, and updates progress.",
            ),
            (
                "server.py",
                "Serves the page and provides the file-backed API.",
                "Handles GET, POST, and DELETE requests for reader records and writes changes to data/readers.json.",
            ),
            (
                "data/readers.json",
                "Stores the leaderboard records.",
                "Keeps currentReaderId and a readers array containing each saved child record.",
            ),
        ],
        [1700, 2500, 5160],
    )

    document.add_heading("How Each File Works", level=1)

    document.add_heading("index.html", level=2)
    document.add_paragraph(
        "The HTML file is the skeleton of the page. It links to styles.css for presentation "
        "and app.js for behavior. The main content is split into a setup panel and a board panel."
    )
    add_bullet(document, "The setup panel contains text fields for name and favorite book.")
    add_bullet(document, "Goal inputs collect daily minutes, weekly minutes, and monthly books.")
    add_bullet(document, "Progress inputs collect today’s minutes, this week’s minutes, and books finished this month.")
    add_bullet(document, "The leaderboard panel includes progress meters and an empty ordered list that app.js fills with records.")

    document.add_heading("styles.css", level=2)
    document.add_paragraph(
        "The CSS file makes the page readable and kid-friendly while keeping the layout practical. "
        "It defines the color tokens, grid layout, form styling, progress meters, leaderboard rows, and mobile behavior."
    )
    add_bullet(document, "The .app-shell grid creates the two-column desktop layout and collapses to one column on smaller screens.")
    add_bullet(document, "The .leader-row style lays out rank, reader details, score, and the Remove button.")
    add_bullet(document, "The .empty-row style shows a friendly placeholder when there are no readers.")
    add_bullet(document, "The mobile media query stacks controls so text and buttons do not overflow.")

    document.add_heading("app.js", level=2)
    document.add_paragraph(
        "The JavaScript file is the browser controller. It never writes the JSON file directly; "
        "instead, it calls the server API and then redraws the page using the returned records."
    )
    add_bullet(document, "loadReaders() calls GET /api/readers and renders every saved reader.")
    add_bullet(document, "addReader() calls POST /api/readers with the form data, adding a new record.")
    add_bullet(document, "removeReader() calls DELETE /api/readers/:id for the clicked row.")
    add_bullet(document, "clearReaders() calls DELETE /api/readers to clear all records.")
    add_bullet(document, "updateLeaderboard() sorts readers by score and creates the visible leaderboard rows.")
    add_bullet(document, "updateProgress() shows goal progress for the most recently added or selected reader.")

    document.add_heading("server.py", level=2)
    document.add_paragraph(
        "The Python server has two jobs: serve the static web files and manage the records file. "
        "It uses the built-in HTTP server, so no extra framework is required."
    )
    add_bullet(document, "GET /api/readers returns the current JSON records.")
    add_bullet(document, "POST /api/readers validates the reader id and name, appends the reader, and saves the file.")
    add_bullet(document, "DELETE /api/readers/:id removes only that reader from the records list.")
    add_bullet(document, "DELETE /api/readers resets the records file to an empty leaderboard.")
    add_bullet(document, "load_records() and save_records() centralize reading and writing data/readers.json.")

    document.add_heading("data/readers.json", level=2)
    document.add_paragraph(
        "This JSON file is the app’s simple database. It is intentionally human-readable, "
        "so you can inspect the saved readers without opening a database tool."
    )
    add_table(
        document,
        ["JSON field", "Meaning"],
        [
            ("currentReaderId", "The id of the most recently added reader, used for rank and progress highlighting."),
            ("readers", "An array of saved reader records displayed on the leaderboard."),
            ("reader.id", "A unique id generated in the browser when a reader is added."),
            ("reader.goals", "The daily, weekly, and monthly goals captured for that reader."),
            ("reader.todayMinutes, reader.weekMinutes, reader.monthBooks", "The progress values used for meters and scoring."),
        ],
        [2300, 7060],
    )

    document.add_heading("Data Flow", level=1)
    document.add_paragraph(
        "The app follows a simple request-and-render loop. The browser sends changes to the server, "
        "the server updates the JSON file, and the browser redraws the leaderboard from the saved response."
    )
    add_table(
        document,
        ["Action", "Browser behavior", "Server/file behavior"],
        [
            (
                "Open page",
                "app.js calls GET /api/readers.",
                "server.py reads data/readers.json and returns all records.",
            ),
            (
                "Add reader",
                "app.js sends the form data with POST /api/readers.",
                "server.py appends the reader to readers.json and returns the full list.",
            ),
            (
                "Remove reader",
                "app.js sends DELETE /api/readers/:id from that row’s button.",
                "server.py filters that id out of readers.json and returns the remaining list.",
            ),
            (
                "Reset",
                "app.js sends DELETE /api/readers.",
                "server.py writes an empty records object to readers.json.",
            ),
        ],
        [1700, 3830, 3830],
    )

    document.add_heading("Running the App", level=1)
    document.add_paragraph("From the project folder, run:")
    command = document.add_paragraph()
    command.paragraph_format.left_indent = Inches(0.25)
    command.paragraph_format.space_after = Pt(6)
    run = command.add_run("python3 server.py")
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    document.add_paragraph("Then open http://127.0.0.1:5173/ in the browser.")

    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
